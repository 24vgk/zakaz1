import io
import os
from collections import defaultdict
from docxtpl import DocxTemplate
from datetime import date
from pathlib import Path
from aiogram import Router, F
from aiogram.exceptions import TelegramBadRequest
from aiogram.types import Message, CallbackQuery, BufferedInputFile, InlineKeyboardButton, InlineKeyboardMarkup, \
    FSInputFile
from aiogram.fsm.state import StatesGroup, State
from aiogram.fsm.context import FSMContext
from aiogram.filters import BaseFilter
from sqlalchemy import select, func, or_, delete
from docx import Document
from sqlalchemy.ext.asyncio import AsyncSession

from config import GROUP_CHAT_ID
from db import session_scope
from crud import upsert_problems, set_report_status, problems_stats, set_admin, set_problem_status, \
    close_list_if_completed, upsert_staff, all_regular_approved, get_admin_ids, split_admins, upsert_review, \
    has_any_rejection
from keyboards.admin_kb import review_kb
from models import ReportStatus, Report, ProblemStatus, Problem, ProblemList, Role, User, Staff, ActEntry, \
    ReportDecision, ReportReview
from utils.parsing import parse_problems_csv, parse_problems_xlsx

from keyboards.admin_main_kb import admin_main_menu
from keyboards.admin_manage_kb import admins_menu, cancel_kb
import matplotlib

from utils.staff_import import parse_staff_xlsx

matplotlib.use("Agg")  # без GUI
import matplotlib.pyplot as plt

class AdminOnly(BaseFilter):
    async def __call__(self, event, **data):
        return data.get("event_from_user_role") == "admin"

admin_router = Router(name="admin")


async def guard_admin(call_or_msg, event_from_user_role: str | None) -> bool:
    if event_from_user_role != "admin":
        # для CallbackQuery и Message поведение одинаковое
        text = "Недостаточно прав. Это действие доступно только администраторам."
        if hasattr(call_or_msg, "answer") and call_or_msg.__class__.__name__ == "CallbackQuery":
            await call_or_msg.answer(text, show_alert=True)
        else:
            await call_or_msg.answer(text)
        return False
    return True


async def build_votes_summary(session: AsyncSession, report_id: int) -> str:
    """
    Возвращает текст вида:
      👥 Голоса админов:
      - Иван (обычный): ✅ Принял
      - Пётр (главный): ⏳ Нет голоса
    """
    all_admin_ids = await get_admin_ids(session)
    regular_ids, main_ids = split_admins(all_admin_ids)

    if not all_admin_ids:
        return ""

    # грузим всех админов
    rows = await session.execute(
        select(User).where(User.id.in_(all_admin_ids))
    )
    admins: list[User] = list(rows.scalars().all())
    admin_by_id = {u.id: u for u in admins}

    # грузим все решения по отчёту
    rows2 = await session.execute(
        select(ReportReview).where(ReportReview.report_id == report_id)
    )
    reviews: list[ReportReview] = list(rows2.scalars().all())
    decision_by_admin: dict[int, ReportDecision] = {
        r.admin_id: r.decision for r in reviews
    }

    # имя пользователя
    def short_name(u: User) -> str:
        if u.first_name or u.last_name:
            return f"{u.first_name or ''} {u.last_name or ''}".strip()
        if u.username:
            return f"@{u.username}"
        return str(u.id)

    lines: list[str] = ["\n👥 Голоса админов:"]

    for aid in all_admin_ids:
        u = admin_by_id.get(aid)
        if not u:
            continue

        role_label = "главный" if aid in main_ids else "обычный"
        dec = decision_by_admin.get(aid)

        if dec == ReportDecision.APPROVED:
            mark = "✅ Принял"
        elif dec == ReportDecision.REJECTED:
            mark = "❌ Отклонил"
        else:
            mark = "⏳ Нет голоса"

        lines.append(f"- {short_name(u)} ({role_label}): {mark}")

    return "\n".join(lines)

class AdminStates(StatesGroup):
    waiting_list_code = State()
    waiting_list_title = State()  # <<< НОВОЕ
    waiting_csv = State()
    waiting_reject_reason = State()
    waiting_add_admin_id = State()
    waiting_del_admin_id = State()
    waiting_staff_file = State()

# ===== Главная админ-панель =====
@admin_router.callback_query(F.data == "admin:back_main")
async def cb_back_main(call: CallbackQuery, state: FSMContext):
    await state.clear()
    await call.message.edit_text("👋 Привет, администратор! Выберите действие:", reply_markup=admin_main_menu())
    await call.answer()


# ===== Создание тестового акта =====
def _docx_replace_all(doc: Document, mapping: dict[str, str]) -> None:
    """Грубая замена {{placeholders}} по всему документу."""
    def _replace_in_run(run, mapping):
        text = run.text
        changed = False
        for k, v in mapping.items():
            placeholder = f"{{{{{k}}}}}"   # {{title}}
            if placeholder in text:
                text = text.replace(placeholder, v)
                changed = True
        if changed:
            run.text = text

    # параграфы
    for p in doc.paragraphs:
        for r in p.runs:
            _replace_in_run(r, mapping)

    # таблицы
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for r in p.runs:
                        _replace_in_run(r, mapping)


@admin_router.callback_query(F.data == "admin:akt")
async def cb_admin_create_akt_by_staff(
    call: CallbackQuery,
    event_from_user_role: str | None = None,
):
    """
    Формирует акты по сотрудникам из таблицы Staff.

    Для каждого staff.assignee:
      - ищем проблемы со статусом ACCEPTED
      - у которых assignees_raw начинается с этого ID (0-й индекс)
      - и по которым ЕЩЁ НЕТ записи в Acts (ActEntry)
      - подгружаем ProblemList (code, title)
      - собираем один акт на сотрудника с его задачами
      - после генерации записываем ActEntry, чтобы второй раз не брать.
    """
    if not await guard_admin(call, event_from_user_role):
        return

    await call.answer("Генерирую акты по исполнителям...", show_alert=False)

    # папка для временных файлов
    os.makedirs("temp", exist_ok=True)
    doc_path = "shablon/akt.docx"

    total_acts = 0

    async with session_scope() as s:
        # 1) Берём всех сотрудников
        staff_rows = (
            await s.execute(select(Staff).order_by(Staff.fio))
        ).scalars().all()

        for staff in staff_rows:
            tg_id = staff.assignee

            # --- подзапрос: есть ли уже акт по этой задаче и этому исполнителю
            act_exists = (
                select(ActEntry.id)
                .where(
                    ActEntry.problem_id == Problem.id,
                    ActEntry.assignee == tg_id,
                )
                .exists()
            )

            # 2) Ищем задачи, где этот tg_id стоит ПЕРВЫМ в assignees_raw
            #    и статус == ACCEPTED
            #    и для них ещё нет записи в Acts
            stmt = (
                select(Problem, ProblemList)
                .join(ProblemList, Problem.list_id == ProblemList.id)
                .where(
                    Problem.status == ProblemStatus.ACCEPTED,
                    Problem.assignees_raw.isnot(None),
                    or_(
                        Problem.assignees_raw == str(tg_id),
                        Problem.assignees_raw.like(f"{tg_id},%"),
                    ),
                    ~act_exists,  # <<< акт ещё не формировался
                )
                .order_by(ProblemList.code, Problem.number)
            )
            rows = (await s.execute(stmt)).all()

            if not rows:
                continue  # у этого сотрудника нет новых принятых задач – пропускаем

            # 3) Собираем текст для {{data}}
            lines: list[str] = []
            for prob, plist in rows:
                lines.append(
                    f"№{prob.number}"
                )
            data_text = ", ".join(lines)

            # 4) Берём данные ProblemList (из первого списка в выборке)
            first_plist: ProblemList = rows[0][1]
            list_title = first_plist.title or first_plist.code
            list_code = first_plist.code

            # 4) Готовим контекст для шаблона
            context = {
                "title": list_title,       # подгони под свой шаблон
                "data": data_text,
                "post": staff.post,
                "fio": staff.fio,
            }

            # 5) Рендерим docx по шаблону
            try:
                doc = DocxTemplate(doc_path)
            except Exception as e:
                await call.message.answer(
                    f"❌ Не удалось открыть шаблон акта: {e}",
                    reply_markup=admin_main_menu(),
                )
                return

            doc.render(context)

            # имя файла: akt_<fio_or_id>.docx
            safe_fio = (staff.fio or str(tg_id)).replace(" ", "_")
            filename = f"akt_{list_code}_{safe_fio}.docx"
            out_path = os.path.join("temp", filename)

            doc.save(out_path)
            total_acts += 1

            # 6) Запоминаем, что по этим задачам и этому исполнителю акт уже сформирован
            for prob, _plist in rows:
                s.add(
                    ActEntry(
                        problem_id=prob.id,
                        assignee=tg_id,
                    )
                )

            # можно коммитить пачками, но одного в конце контекста обычно достаточно.
            # я добавлю явный коммит после цикла по staff.
            await call.message.answer_document(
                document=FSInputFile(out_path),
                caption=f"Акт для {staff.fio or tg_id}",
            )
            await call.message.answer_document(
                document=FSInputFile(out_path),
                caption=f"Акт для {staff.fio or tg_id}",
            )

        # фиксируем все ActEntry
        await s.commit()

    # 7) Итоговое сообщение
    if total_acts == 0:
        await call.message.answer(
            "Не найдено новых задач в статусе <b>Принято</b> для формирования актов.",
            reply_markup=admin_main_menu(),
        )
    else:
        await call.message.answer(
            f"Готово! Сформировано актов: <b>{total_acts}</b>.",
            reply_markup=admin_main_menu(),
        )

# @admin_router.callback_query(F.data == "admin:akt")
# async def cb_admin_create_akt(call: CallbackQuery):
#     await call.answer("Генерирую акты...", show_alert=False)
#
#     # Создаём папку для актов, если её нет
#     os.makedirs("temp", exist_ok=True)
#
#     doc_path = "shablon/akt.docx"
#     generated_files = []
#
#     async with session_scope() as s:
#         # Получаем все списки
#         result = await s.execute(select(ProblemList))
#         lists = result.scalars().all()
#
#         for plist in lists:
#             # Берём все решённые проблемы
#             res = await s.execute(
#                 select(Problem)
#                 .where(
#                     Problem.list_id == plist.id,
#                     Problem.status == ProblemStatus.ACCEPTED
#                 )
#             )
#             problems = res.scalars().all()
#
#             if not problems:
#                 continue
#
#             # Группируем по ответственным
#             grouped: dict[int, list[Problem]] = {}
#
#             for prob in problems:
#                 for tg_id in prob.assignees:
#                     grouped.setdefault(tg_id, []).append(prob)
#
#             for tg_id, probs in grouped.items():
#                 # Берём сотрудника
#                 st = await s.execute(
#                     select(Staff).where(Staff.assignee == tg_id)
#                 )
#                 staff = st.scalar_one_or_none()
#
#                 if not staff:
#                     continue
#
#                 # Формируем список задач
#                 data_text = ", ".join(f"№{prob.number}" for prob in probs)
#
#                 context = {
#                     "title": plist.title,
#                     "data": data_text,
#                     "post": staff.post,
#                     "fio": staff.fio,
#                 }
#
#                 doc = DocxTemplate(doc_path)
#                 doc.render(context)
#
#                 safe_code = plist.code.replace(" ", "_")
#                 out_name = f"akt_{safe_code}_{tg_id}.docx"
#                 out_path = os.path.join("temp", out_name)
#
#                 doc.save(out_path)
#                 generated_files.append(out_path)
#
#     # Отправляем акты
#     for path in generated_files:
#         await call.message.answer_document(FSInputFile(path))
#
#     if not generated_files:
#         await call.message.answer("Нет списков с решёнными задачами — акты не созданы.")
#     else:
#         await call.message.answer(f"Готово! Создано актов: {len(generated_files)}")

# ===== Загрузка работников =====


@admin_router.callback_query(F.data == "admin:upload_staff")
async def cb_admin_upload_staff(
    call: CallbackQuery,
    state: FSMContext,
    event_from_user_role: str | None = None,
):
    if not await guard_admin(call, event_from_user_role):
        return

    await state.set_state(AdminStates.waiting_staff_file)
    await call.message.edit_text(
        "Пришлите Excel-файл (.xlsx) со списком сотрудников.\n\n"
        "Ожидаемые колонки:\n"
        "• assignee — Telegram ID\n"
        "• post — должность\n"
        "• fio — ФИО",
        reply_markup=admin_main_menu(),
    )
    await call.answer()


@admin_router.message(AdminStates.waiting_staff_file)
async def msg_admin_staff_file(
    msg: Message,
    state: FSMContext,
    event_from_user_role: str | None = None,
):
    # защита по роли
    if not await guard_admin(msg, event_from_user_role):
        await state.clear()
        return

    # проверяем, что это документ
    if not msg.document:
        await msg.answer("Пожалуйста, пришлите файл в формате .xlsx как документ.")
        return

    filename = msg.document.file_name or ""
    if not filename.lower().endswith(".xlsx"):
        await msg.answer("Нужен файл в формате .xlsx.")
        return

    try:
        file = await msg.bot.get_file(msg.document.file_id)
        raw = await msg.bot.download_file(file.file_path)
        data = raw.read()
    except Exception as e:
        await msg.answer(f"❌ Не удалось скачать файл.\nОшибка: {e}")
        await state.clear()
        return

    # парсим
    try:
        rows = parse_staff_xlsx(data)
        if not rows:
            await msg.answer("Файл прочитан, но не найдено ни одной корректной строки.")
            await state.clear()
            return
    except Exception as e:
        await msg.answer(f"❌ Не удалось прочитать файл.\nОшибка: {e}")
        await state.clear()
        return

    # пишем в БД
    try:
        async with session_scope() as s:
            count = await upsert_staff(s, rows)
    except Exception as e:
        await msg.answer(f"❌ Не удалось обновить данные сотрудников.\nОшибка: {e}")
        await state.clear()
        return

    await msg.answer(
        f"✅ Данные сотрудников обновлены.\nОбработано записей: {count}.",
        reply_markup=admin_main_menu(),
    )
    await state.clear()

# ===== Загрузка проблем (кнопка) =====


@admin_router.callback_query(F.data == "admin:upload_problems")
async def cb_admin_upload(call: CallbackQuery, state: FSMContext, event_from_user_role: str | None = None):
    await state.set_state(AdminStates.waiting_list_code)
    await call.message.edit_text(
        "Введите название списка проблем:",
        reply_markup=cancel_kb()
    )
    await call.answer()

@admin_router.message(AdminStates.waiting_list_code)
async def receive_list_code(msg: Message, state: FSMContext, event_from_user_role: str | None = None):
    code = (msg.text or "").strip()
    if not code:
        await msg.answer("Название списка не должено быть пустым.", reply_markup=cancel_kb())
        return

    await state.update_data(list_code=code)

    # ТЕПЕРЬ спрашиваем название списка
    await state.set_state(AdminStates.waiting_list_title)
    await msg.answer(
        f"✔ Наименование: <b>{code}</b>\n\nВведите номер Акта в формате '№10 от 20.10.2025':",
        reply_markup=cancel_kb()
    )

@admin_router.message(AdminStates.waiting_list_title)
async def receive_list_title(msg: Message, state: FSMContext, event_from_user_role: str | None = None):
    title = (msg.text or "").strip()
    if not title:
        await msg.answer("Номер Акта не должен быть пустым.", reply_markup=cancel_kb())
        return

    await state.update_data(list_title=title)

    await state.set_state(AdminStates.waiting_csv)
    await msg.answer(
        f"Номер Акта: <b>{title}</b>\nТеперь отправьте XLSX файл с задачами:",
        reply_markup=cancel_kb()
    )

@admin_router.message(AdminStates.waiting_csv, F.document)
async def handle_table(msg: Message, state: FSMContext, event_from_user_role: str | None = None):
    if not await guard_admin(msg, event_from_user_role):
        return

    try:
        file = await msg.bot.get_file(msg.document.file_id)
        buf = await msg.bot.download_file(file.file_path)
        data = buf.read()
        name = (msg.document.file_name or "").lower()

        # код списка = имя файла без расширения
        data_state = await state.get_data()
        list_code = data_state.get("list_code")  # <- берём введённое админом имя
        list_title = data_state.get("list_title")
        list_code_file = Path(msg.document.file_name or "problems").stem

        # разбираем файл по твоему шаблону
        if name.endswith(".xlsx"):
            rows = list(parse_problems_xlsx(data))
        else:
            raise ValueError("Ожидается .xlsx с колонками: id, title, assignee, due_date")

        if not rows:
            await msg.answer("Файл прочитан, но в нём нет ни одной строки с задачами.", reply_markup=admin_main_menu())
            await state.clear()
            return

        # обновляем/создаём список и его проблемы
        async with session_scope() as s:
            plist = await upsert_problems(s, list_code, rows, list_title=list_title)

        # создаём тему в группе для этого списка (если указана GROUP_CHAT_ID)
        if GROUP_CHAT_ID:
            async with session_scope() as s:
                # перечитаем список уже с сессией
                result = await s.execute(
                    select(ProblemList).where(ProblemList.code == list_code)
                )
                plist_db = result.scalar_one_or_none()
                if plist_db and plist_db.group_topic_id is None:
                    try:
                        topic = await msg.bot.create_forum_topic(
                            chat_id=GROUP_CHAT_ID,
                            name=plist_db.title or plist_db.code,
                        )
                        plist_db.group_topic_id = topic.message_thread_id
                        await s.commit()
                    except Exception as e:
                        # не валим бота, если нет прав / группа без тем и т.п.
                        print(f"Не удалось создать тему для списка {list_code}: {e}")

        await msg.answer(
            f"✅ Список проблем '{list_code}' загружен.\n"
            f"Задач в файле: {len(rows)}.\n"
            f"Тема в группе {'создана' if GROUP_CHAT_ID else 'не настраивалась (нет GROUP_CHAT_ID)'}",
            reply_markup=admin_main_menu(),
        )
        await state.clear()

    except Exception as e:
        await msg.answer(
            f"❌ Не удалось прочитать файл или обновить данные.\nОшибка: {e}",
            reply_markup=admin_main_menu(),
        )
        await state.clear()


@admin_router.callback_query(F.data == "admin:cancel")
async def cb_cancel(call: CallbackQuery, state: FSMContext):
    await state.clear()
    await call.message.edit_text("Действие отменено.", reply_markup=admin_main_menu())
    await call.answer()

async def _send_list_stats(message, list_code: str):
    """
    Рисует круговую диаграмму по ВСЕМ проблемам списка list_code
    и отправляет её как фото.
    Показывает 4 статуса: В работе, Отправлен отчёт, Принято, Отклонено.
    Отдельно показывает количество просроченных (непринятых) задач.
    """
    today_str = date.today().strftime("%Y-%m-%d")

    async with session_scope() as s:
        # --- агрегация по статусам ---
        rows = await s.execute(
            select(
                Problem.status,
                func.count(Problem.id)
            )
            .join(ProblemList)
            .where(ProblemList.code == list_code)
            .group_by(Problem.status)
        )
        rows = rows.all()

        # --- количество просроченных задач ---
        overdue_q = await s.execute(
            select(func.count(Problem.id))
            .join(ProblemList)
            .where(
                ProblemList.code == list_code,
                Problem.status != ProblemStatus.ACCEPTED,
                Problem.due_date.isnot(None),
                Problem.due_date < today_str,   # 'YYYY-MM-DD' – строковое сравнение корректно
            )
        )
        overdue_total = overdue_q.scalar_one() or 0

    if not rows:
        await message.answer(
            f"В списке <b>{list_code}</b> нет проблем.",
            reply_markup=admin_main_menu(),
        )
        return

    # --- раскладываем по статусам ---
    by_status: dict[ProblemStatus, int] = {st: cnt for st, cnt in rows}

    in_work     = by_status.get(ProblemStatus.IN_PROGRESS, 0)
    report_sent = by_status.get(ProblemStatus.REPORT_SENT, 0)
    accepted    = by_status.get(ProblemStatus.ACCEPTED, 0)
    rejected    = by_status.get(ProblemStatus.REJECTED, 0)

    total = in_work + report_sent + accepted + rejected

    if total == 0:
        await message.answer(
            f"В списке <b>{list_code}</b> нет активных проблем.",
            reply_markup=admin_main_menu(),
        )
        return

    # --- готовим данные для диаграммы ---
    labels: list[str] = []
    sizes: list[int] = []
    colors: list[str] = []

    if in_work > 0:
        labels.append("В работе")
        sizes.append(in_work)
        colors.append("#FFD700")  # 🟡

    if report_sent > 0:
        labels.append("Отправлен отчёт")
        sizes.append(report_sent)
        colors.append("#1E90FF")  # 🔵

    if accepted > 0:
        labels.append("Принято")
        sizes.append(accepted)
        colors.append("#32CD32")  # 🟢

    if rejected > 0:
        labels.append("Отклонено")
        sizes.append(rejected)
        colors.append("#FF4500")  # 🔴

    # --- рисуем круговую диаграмму ---
    fig, ax = plt.subplots(figsize=(5, 5))

    explode = [0.05] * len(sizes)

    wedges, texts, autotexts = ax.pie(
        sizes,
        labels=labels,
        colors=colors,  # цвета синхронизированы с caption
        autopct=lambda pct: f"{pct:.1f}%",
        explode=explode,
        startangle=90,
        shadow=True,
    )

    for autot in autotexts:
        autot.set_size(9)

    ax.set_title(f"Статистика по списку {list_code}")
    ax.axis("equal")
    plt.tight_layout()

    # --- сохраняем в буфер ---
    buf = io.BytesIO()
    plt.savefig(buf, format="png")
    plt.close(fig)
    buf.seek(0)

    photo = BufferedInputFile(buf.getvalue(), filename=f"stats_{list_code}.png")

    # --- текстовая сводка под картинкой ---
    caption = (
        f"📊 <b>Статистика по списку {list_code}</b>\n\n"
        f"Всего проблем: {total}\n"
        f"⏰ Просрочено (не принято): {overdue_total}\n\n"
        f"🟡 В работе: {in_work}\n"
        f"🔵 Отправлен отчёт: {report_sent}\n"
        f"🟢 Принято: {accepted}\n"
        f"🔴 Отклонено: {rejected}"
    )

    await message.answer_photo(
        photo=photo,
        caption=caption,
        reply_markup=admin_main_menu(),
    )

@admin_router.callback_query(F.data == "admin:stats_problems")
async def cb_admin_stats(call: CallbackQuery, event_from_user_role: str | None = None):
    if not await guard_admin(call, event_from_user_role):
        return

    # найдём все списки, в которых есть проблемы
    async with session_scope() as s:
        rows = await s.execute(
            select(ProblemList.code)
            .join(Problem, Problem.list_id == ProblemList.id)
            .group_by(ProblemList.code)
            .order_by(ProblemList.code)
        )
        codes = [r[0] for r in rows.all()]

    if not codes:
        await call.message.edit_text(
            "Нет ни одного списка проблем.",
            reply_markup=admin_main_menu(),
        )
        await call.answer()
        return

    # если список только один — сразу рисуем по нему круг
    if len(codes) == 1:
        await _send_list_stats(call.message, codes[0])
        await call.answer()
        return

    # иначе даём выбор списка
    kb = InlineKeyboardMarkup(
        inline_keyboard=[
            [InlineKeyboardButton(text=code, callback_data=f"admin:stats_problems_list:{code}")]
            for code in codes
        ] + [
            [InlineKeyboardButton(text="⬅️ Назад", callback_data="admin:back_main")]
        ]
    )

    await call.message.edit_text(
        "Выберите список для статистики:",
        reply_markup=kb,
    )
    await call.answer()


@admin_router.callback_query(F.data.startswith("admin:stats_problems_list:"))
async def cb_admin_stats_list(call: CallbackQuery, event_from_user_role: str | None = None):
    if not await guard_admin(call, event_from_user_role):
        return
    _, _, list_code = call.data.split(":", 3)
    await _send_list_stats(call.message, list_code)
    await call.answer()

# ===== Управление администраторами (кнопка -> подменю) =====
@admin_router.callback_query(F.data == "admin:admins")
async def cb_admins_menu(call: CallbackQuery, state: FSMContext, event_from_user_role: str | None = None):
    if not await guard_admin(call, event_from_user_role):
        return
    await state.clear()
    await call.message.edit_text("Выберите действие:", reply_markup=admins_menu())
    await call.answer()

@admin_router.callback_query(F.data == "admin:add_admin")
async def cb_add_admin_prompt(call: CallbackQuery, state: FSMContext):
    await state.set_state(AdminStates.waiting_add_admin_id)
    await call.message.edit_text(
        "Введите числовой Telegram ID пользователя, которого нужно сделать администратором:",
        reply_markup=cancel_kb()
    )
    await call.answer()

@admin_router.message(AdminStates.waiting_add_admin_id)
async def add_admin_receive_id(msg: Message, state: FSMContext):
    if not msg.text or not msg.text.isdigit():
        await msg.answer("Нужны только цифры. Попробуйте снова.", reply_markup=cancel_kb())
        return
    target = int(msg.text)
    async with session_scope() as s:
        await set_admin(s, target, True)
    await msg.answer(f"✅ Пользователь {target} теперь администратор.", reply_markup=admins_menu())
    await state.clear()

@admin_router.callback_query(F.data == "admin:del_admin")
async def cb_del_admin_prompt(call: CallbackQuery, state: FSMContext, event_from_user_role: str | None = None):
    if not await guard_admin(call, event_from_user_role):
        return
    await state.set_state(AdminStates.waiting_del_admin_id)
    await call.message.edit_text(
        "Введите числовой Telegram ID пользователя, у которого нужно забрать права администратора:",
        reply_markup=cancel_kb()
    )
    await call.answer()

@admin_router.message(AdminStates.waiting_del_admin_id)
async def del_admin_receive_id(msg: Message, state: FSMContext):
    if not msg.text or not msg.text.isdigit():
        await msg.answer("Нужны только цифры. Попробуйте снова.", reply_markup=cancel_kb())
        return
    target = int(msg.text)
    async with session_scope() as s:
        await set_admin(s, target, False)
    await msg.answer(f"✅ Пользователь {target} теперь пользователь.", reply_markup=admins_menu())
    await state.clear()

# ===== Модерация отчётов (кнопки уже были) =====
@admin_router.callback_query(F.data.startswith("admin:accept:"))
async def cb_accept(
    call: CallbackQuery,
    event_from_user_role: str | None = None,
):
    if not await guard_admin(call, event_from_user_role):
        return

    try:
        _, _, report_id_s, user_tg_s = call.data.split(":", 3)
        report_id = int(report_id_s)
        user_tg_id = int(user_tg_s)
    except Exception:
        await call.answer("Некорректные данные кнопки.", show_alert=True)
        return

    admin_tg_id = call.from_user.id

    async with session_scope() as s:
        report = await s.get(Report, report_id)
        if not report:
            await call.answer("Отчёт не найден.", show_alert=True)
            return

        problem = await s.get(Problem, report.problem_id)

        all_admin_ids = await get_admin_ids(s)
        regular_ids, main_ids = split_admins(all_admin_ids)

        # фиксируем, что этот админ одобрил
        await upsert_review(s, report_id, admin_tg_id, ReportDecision.APPROVED)

        # если уже кто-то отклонил — не даём принять
        if await has_any_rejection(s, report_id):
            await s.commit()
            await call.answer("По этому отчёту уже есть отклонение.", show_alert=True)
            return

        # === если это главный (этап 2/2) ===
        if admin_tg_id in main_ids:
            report.status = ReportStatus.ACCEPTED
            if problem:
                problem.status = ProblemStatus.ACCEPTED
                problem.note = None
            report.admin_id = admin_tg_id

            await s.commit()

            # юзеру
            try:
                await call.bot.send_message(
                    chat_id=user_tg_id,
                    text="✅ Ваш отчёт по задаче принят.",
                )
            except Exception:
                pass

            # собираем сводку голосов
            votes_text = await build_votes_summary(s, report_id)

            # убираем кнопки на этом сообщении и дописываем статус + голоса
            try:
                base = call.message.caption or call.message.text or ""
                new_text = base + "\n\n✅ Окончательно принято." + votes_text
                if call.message.caption is not None:
                    await call.message.edit_caption(new_text, reply_markup=None)
                else:
                    await call.message.edit_text(new_text, reply_markup=None)
            except Exception:
                pass

            await call.answer("Отчёт окончательно принят ✅")
            return

        # === обычный админ (этап 1/2) ===
        # если все обычные одобрили — шлём главным
        if await all_regular_approved(s, report_id, regular_ids):
            for mid in main_ids:
                try:
                    await call.bot.copy_message(
                        chat_id=mid,
                        from_chat_id=report.user_chat_id,
                        message_id=report.user_msg_id,
                        caption=(
                            f"Новый отчёт #{report.id} (этап 2/2)\n"
                            f"Все обычные админы его одобрили.\n\n"
                            f"Нажмите, чтобы принять или отклонить."
                        ),
                        reply_markup=review_kb(report.id, user_tg_id),
                    )
                except Exception:
                    pass

        await s.commit()

        # обновляем текст у ТЕКУЩЕГО админа (его копии) — статус + голоса
        try:
            votes_text = await build_votes_summary(s, report_id)
            suffix = "\n\n✅ Ваш голос 'Принять' учтён." + votes_text
            base = call.message.caption or call.message.text or ""
            new_text = base + suffix

            if call.message.caption is not None:
                await call.message.edit_caption(new_text, reply_markup=None)
            else:
                await call.message.edit_text(new_text, reply_markup=None)
        except Exception:
            pass

    await call.answer("Ваше одобрение учтено ✅")

@admin_router.callback_query(F.data.startswith("admin:reject:"))
async def cb_reject_start(
    call: CallbackQuery,
    state: FSMContext,
    event_from_user_role: str | None = None,
):
    if not await guard_admin(call, event_from_user_role):
        return

    try:
        _, _, report_id_s, user_tg_s = call.data.split(":", 3)
        report_id = int(report_id_s)
        user_tg_id = int(user_tg_s)
    except Exception:
        await call.answer("Некорректные данные кнопки.", show_alert=True)
        return

    # сохраняем контекст в FSM
    await state.update_data(
        reject_report_id=report_id,
        reject_user_tg_id=user_tg_id,
        reject_message_chat_id=call.message.chat.id,
        reject_message_id=call.message.message_id,
    )
    await state.set_state(AdminStates.waiting_reject_reason)

    await call.message.answer(
        "❌ Вы собираетесь отклонить отчёт.\n"
        "Пожалуйста, введите причину отклонения одним сообщением."
    )
    await call.answer()

@admin_router.message(AdminStates.waiting_reject_reason)
async def cb_reject_reason(
    msg: Message,
    state: FSMContext,
    event_from_user_role: str | None = None,
):
    if not await guard_admin(msg, event_from_user_role):
        await state.clear()
        return

    data = await state.get_data()
    report_id = int(data.get("reject_report_id"))
    user_tg_id = int(data.get("reject_user_tg_id"))
    msg_chat_id = data.get("reject_message_chat_id")
    msg_id = data.get("reject_message_id")

    reason = (msg.text or "").strip()
    if not reason:
        await msg.answer("Причина не может быть пустой. Попробуйте ещё раз или /cancel.")
        return

    admin_tg_id = msg.from_user.id

    async with session_scope() as s:
        report = await s.get(Report, report_id)
        if not report:
            await msg.answer("Отчёт не найден.", reply_markup=admin_main_menu())
            await state.clear()
            return

        problem = await s.get(Problem, report.problem_id)

        # фиксируем REJECT
        await upsert_review(s, report_id, admin_tg_id, ReportDecision.REJECTED)

        # отчёт и задача сразу в REJECTED
        report.status = ReportStatus.REJECTED
        report.admin_id = admin_tg_id
        report.admin_reason = reason

        if problem:
            problem.status = ProblemStatus.REJECTED
            problem.note = reason

        await s.commit()

        # уведомляем исполнителя
        try:
            text = (
                "❌ Ваш отчёт отклонён. Задача отправлена на доработку.\n"
                f"Причина: {reason}"
            )
            await msg.bot.send_message(
                chat_id=user_tg_id,
                text=text,
            )
        except Exception:
            pass

        # для красоты — получим сводку голосов
        votes_text = await build_votes_summary(s, report_id)

    # обновляем подпись/текст у исходного сообщения с отчётом (у этого админа)
    try:
        if msg_chat_id and msg_id:
            # подтягиваем сообщение через bot.edit...
            from aiogram.exceptions import TelegramBadRequest

            try:
                await msg.bot.edit_message_caption(
                    chat_id=msg_chat_id,
                    message_id=msg_id,
                    caption=(
                        f"Отчёт #{report_id}\n\n"
                        f"❌ Отклонён. Задача на доработку.\n"
                        f"Причина: {reason}"
                        f"{votes_text}"
                    ),
                    reply_markup=None,
                )
            except TelegramBadRequest:
                # если не было подписи — пробуем текстом
                await msg.bot.edit_message_text(
                    chat_id=msg_chat_id,
                    message_id=msg_id,
                    text=(
                        f"Отчёт #{report_id}\n\n"
                        f"❌ Отклонён. Задача на доработку.\n"
                        f"Причина: {reason}"
                        f"{votes_text}"
                    ),
                    reply_markup=None,
                )
    except Exception:
        pass

    await msg.answer("❌ Отчёт отклонён. Причина сохранена.", reply_markup=admin_main_menu())
    await state.clear()


@admin_router.callback_query(F.data == "admin:users")
async def cb_admin_users(call: CallbackQuery, event_from_user_role: str | None = None):
    if not await guard_admin(call, event_from_user_role):
        return

    async with session_scope() as s:
        res = await s.execute(select(User).order_by(User.role, User.id))
        users = res.scalars().all()

    if not users:
        # тут тоже безопаснее отвечать новым сообщением
        await call.message.answer(
            "Пользователей в БД пока нет.",
            reply_markup=admin_main_menu(),
        )
        await call.answer()
        return

    admins = [u for u in users if u.role == Role.ADMIN]
    regular = [u for u in users if u.role == Role.USER]

    def fmt_user(u: User) -> str:
        name = " ".join(filter(None, [u.first_name, u.last_name])).strip()
        if not name:
            name = u.username or ""
        return f"{u.id} - {name or 'без имени'} - {u.role.value}"

    lines: list[str] = []

    if admins:
        lines.append("<b>Администраторы:</b>")
        lines += [f"• {fmt_user(u)}" for u in admins]
        lines.append("")

    if regular:
        lines.append("<b>Пользователи:</b>")
        lines += [f"• {fmt_user(u)}" for u in regular]

    text = "\n".join(lines)

    # Кнопки для "копирования" ID
    kb_rows = []
    for u in users[:50]:
        label_name = u.first_name or u.username or "user"
        kb_rows.append([
            InlineKeyboardButton(
                text=f"{label_name} ({u.id})",
                callback_data=f"admin:userid:{u.id}",
            )
        ])
    kb_rows.append([InlineKeyboardButton(text="⬅️ Назад", callback_data="admin:back_main")])
    kb = InlineKeyboardMarkup(inline_keyboard=kb_rows)

    # 🔧 главное изменение:
    if call.message.text:
        # если это обычное текстовое сообщение — редактируем
        await call.message.edit_text(text, reply_markup=kb)
    else:
        # если это медиа / что-то без текста — шлём новое
        await call.message.answer(text, reply_markup=kb)

    await call.answer()


# ===== Удаление списков проблем =====

@admin_router.callback_query(F.data == "admin:delete_plists")
async def cb_admin_delete_plists(
    call: CallbackQuery,
    event_from_user_role: str | None = None,
):
    """Показать список всех списков проблем для выбора удаления."""
    if not await guard_admin(call, event_from_user_role):
        return

    async with session_scope() as s:
        rows = await s.execute(
            select(
                ProblemList.id,
                ProblemList.code,
                ProblemList.title,
                ProblemList.is_closed,
                func.count(Problem.id).label("cnt"),
            )
            .join(Problem, Problem.list_id == ProblemList.id, isouter=True)
            .group_by(ProblemList.id, ProblemList.code, ProblemList.title, ProblemList.is_closed)
            .order_by(ProblemList.code)
        )
        items = rows.all()

    if not items:
        await call.message.edit_text(
            "Списков проблем пока нет.",
            reply_markup=admin_main_menu(),
        )
        await call.answer()
        return

    kb_rows: list[list[InlineKeyboardButton]] = []
    for pid, code, title, is_closed, cnt in items:
        title_display = title or "(без названия)"
        status = "✅ закрыт" if is_closed else "🟢 открыт"
        text = f"{code} — {title_display} ({cnt} задач, {status})"
        kb_rows.append([
            InlineKeyboardButton(
                text=text,
                callback_data=f"admin:del_plist:{code}",
            )
        ])

    kb_rows.append([
        InlineKeyboardButton(
            text="⬅️ Назад в меню",
            callback_data="admin:back_main",
        )
    ])

    kb = InlineKeyboardMarkup(inline_keyboard=kb_rows)

    await call.message.edit_text(
        "Выберите список проблем, который нужно удалить.\n"
        "⚠️ <b>Вместе со списком будут удалены все его задачи.</b>",
        reply_markup=kb,
    )
    await call.answer()


@admin_router.callback_query(F.data.startswith("admin:del_plist:"))
async def cb_admin_del_plist_confirm(
    call: CallbackQuery,
    event_from_user_role: str | None = None,
):
    """Подтверждение удаления выбранного списка проблем."""
    if not await guard_admin(call, event_from_user_role):
        return

    try:
        _, _, list_code = call.data.split(":", 2)
    except Exception:
        await call.answer("Некорректные данные кнопки.", show_alert=True)
        return

    async with session_scope() as s:
        row = await s.execute(
            select(
                ProblemList,
                func.count(Problem.id).label("cnt"),
            )
            .join(Problem, Problem.list_id == ProblemList.id, isouter=True)
            .where(ProblemList.code == list_code)
            .group_by(ProblemList.id)
        )
        res = row.first()

    if not res:
        await call.message.edit_text(
            f"Список с кодом <b>{list_code}</b> не найден.",
            reply_markup=admin_main_menu(),
        )
        await call.answer()
        return

    plist, cnt = res
    title = plist.title or "(без названия)"
    status = "закрыт" if plist.is_closed else "открыт"

    text = (
        f"Вы собираетесь удалить список проблем:\n\n"
        f"<b>Код:</b> {plist.code}\n"
        f"<b>Название:</b> {title}\n"
        f"<b>Статус:</b> {status}\n"
        f"<b>Количество задач:</b> {cnt}\n\n"
        f"⚠️ <b>Все задачи этого списка будут удалены безвозвратно.</b>\n\n"
        f"Продолжить?"
    )

    kb = InlineKeyboardMarkup(inline_keyboard=[
        [
            InlineKeyboardButton(
                text="✅ Да, удалить",
                callback_data=f"admin:del_plist_do:{plist.code}",
            )
        ],
        [
            InlineKeyboardButton(
                text="⬅️ Отмена",
                callback_data="admin:delete_plists",  # вернёмся к выбору списков
            )
        ],
    ])

    await call.message.edit_text(text, reply_markup=kb)
    await call.answer()

@admin_router.callback_query(F.data.startswith("admin:del_plist_do:"))
async def cb_admin_del_plist_do(
    call: CallbackQuery,
    event_from_user_role: str | None = None,
):
    """Удаляет список проблем и все его задачи (без ленивых загрузок)."""
    if not await guard_admin(call, event_from_user_role):
        return

    try:
        _, _, list_code = call.data.split(":", 2)
    except Exception:
        await call.answer("Некорректные данные кнопки.", show_alert=True)
        return

    async with session_scope() as s:
        # 1) найдём список
        row = await s.execute(
            select(ProblemList.id, ProblemList.code, ProblemList.title)
            .where(ProblemList.code == list_code)
        )
        plist_row = row.first()

        if not plist_row:
            await call.message.edit_text(
                f"Список с кодом <b>{list_code}</b> уже не существует.",
                reply_markup=admin_main_menu(),
            )
            await call.answer()
            return

        plist_id, code, title = plist_row
        title = title or "(без названия)"

        # 2) посчитаем количество задач в этом списке
        q_cnt = await s.execute(
            select(func.count(Problem.id)).where(Problem.list_id == plist_id)
        )
        problems_count = q_cnt.scalar_one() or 0

        # 3) сначала удаляем задачи этого списка
        await s.execute(
            delete(Problem).where(Problem.list_id == plist_id)
        )
        # 4) затем удаляем сам список
        await s.execute(
            delete(ProblemList).where(ProblemList.id == plist_id)
        )

        await s.commit()

    text = (
        f"🗑 Список проблем <b>{code}</b> ({title}) удалён.\n"
        f"Удалено задач: {problems_count}."
    )

    try:
        await call.message.edit_text(text, reply_markup=admin_main_menu())
    except TelegramBadRequest as e:
        if "message is not modified" not in str(e):
            raise

    await call.answer("Список удалён ✅", show_alert=False)